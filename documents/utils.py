# documents/utils.py

import os
import pdfplumber
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from docx import Document as DocxDocument
from docx.shared import Pt
from PIL import Image
from django.core.files import File


def convert_document(document):
    """
    Takes a Document model instance.
    Reads document.file and document.conversion_type.
    Produces a converted file and saves it to document.converted_file.
    Returns (success: bool, message: str)
    """

    input_path = document.file.path
    # .path gives the absolute filesystem path to the uploaded file
    # e.g. C:/Users/.../media/documents/myfile.pdf

    conversion = document.conversion_type

    # Build the output file path
    # We strip the original extension and add the new one
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    # os.path.basename gets just the filename: "myfile.pdf"
    # os.path.splitext splits it into ("myfile", ".pdf")
    # [0] gets just "myfile" — the name without extension

    output_dir = os.path.join(os.path.dirname(input_path), "..", "converted")
    # Build path to media/converted/ folder
    # os.path.dirname gets the folder containing the file
    # ".." goes one level up
    # then we go into "converted/"

    os.makedirs(output_dir, exist_ok=True)
    # Create the converted/ folder if it doesn't exist yet
    # exist_ok=True means "don't crash if it already exists"

    try:
        if conversion == "word_to_pdf":
            output_path = _word_to_pdf(input_path, output_dir, base_name)

        elif conversion == "pdf_to_txt":
            output_path = _pdf_to_txt(input_path, output_dir, base_name)

        elif conversion == "pdf_to_word":
            output_path = _pdf_to_word(input_path, output_dir, base_name)

        elif conversion == "txt_to_pdf":
            output_path = _txt_to_pdf(input_path, output_dir, base_name)

        elif conversion == "jpg_to_pdf":
            output_path = _jpg_to_pdf(input_path, output_dir, base_name)

        elif conversion == "pdf_to_jpg":
            output_path = _pdf_to_jpg(input_path, output_dir, base_name)

        elif conversion == "html_to_pdf":
            output_path = _html_to_pdf(input_path, output_dir, base_name)

        else:
            return False, f"Conversion type '{conversion}' is not supported yet."

        # Save the converted file back onto the Document model
        with open(output_path, "rb") as f:
            # "rb" = read binary — always use this for non-text files like PDFs
            output_filename = os.path.basename(output_path)
            document.converted_file.save(output_filename, File(f), save=False)
            # File(f) wraps the Python file object into Django's file wrapper
            # save=False means "don't save the whole model to DB yet"
            # we'll do that ourselves below so we can set other fields too

        document.converted = True
        document.save()
        # Now save everything to the database at once

        return True, "Conversion successful!"

    except Exception as e:
        return False, f"Conversion failed: {str(e)}"
        # str(e) converts the exception into a readable error message


# ── Individual conversion functions ──────────────────────────────────────────

def _word_to_pdf(input_path, output_dir, base_name):
    """Convert DOCX → PDF using python-docx to read + reportlab to write"""

    output_path = os.path.join(output_dir, f"{base_name}.pdf")

    # Read the Word document
    doc = DocxDocument(input_path)

    # Create a PDF using reportlab
    pdf = canvas.Canvas(output_path, pagesize=A4)
    # canvas.Canvas is reportlab's drawing surface
    # pagesize=A4 sets the page dimensions (595 x 842 points)

    width, height = A4
    # A4 is a tuple: (595.27, 841.89) — width and height in points
    # 1 point = 1/72 inch

    y = height - 50
    # Start writing 50 points from the top
    # PDF coordinates start from BOTTOM-LEFT, so higher y = higher on page

    pdf.setFont("Helvetica", 12)
    # Set the font and size for text we're about to draw

    for para in doc.paragraphs:
        # doc.paragraphs is a list of all paragraphs in the Word doc
        text = para.text.strip()

        if not text:
            y -= 15  # blank line — just move down a bit
            continue

        # Handle long lines — wrap them manually
        # reportlab doesn't auto-wrap text, so we do it ourselves
        words = text.split()
        line = ""
        for word in words:
            test_line = f"{line} {word}".strip()
            # pdf.stringWidth measures how wide this text would be in points
            if pdf.stringWidth(test_line, "Helvetica", 12) < width - 80:
                line = test_line
            else:
                pdf.drawString(40, y, line)
                # drawString(x, y, text) — draws text at position x,y
                y -= 18  # move down one line
                line = word

                if y < 50:
                    pdf.showPage()
                    # showPage() ends the current page and starts a new one
                    y = height - 50
                    pdf.setFont("Helvetica", 12)

        if line:
            pdf.drawString(40, y, line)
            y -= 18

        if y < 50:
            pdf.showPage()
            y = height - 50
            pdf.setFont("Helvetica", 12)

    pdf.save()
    # Finalise and write the PDF file to disk
    return output_path


def _pdf_to_txt(input_path, output_dir, base_name):
    """Convert PDF → TXT using pdfplumber"""

    output_path = os.path.join(output_dir, f"{base_name}.txt")

    with pdfplumber.open(input_path) as pdf:
        # pdfplumber.open() reads the PDF
        # 'with' statement ensures the file is properly closed after reading

        all_text = []
        for page in pdf.pages:
            # pdf.pages is a list of page objects
            text = page.extract_text()
            # extract_text() reads all the text from one page
            if text:
                all_text.append(text)

    with open(output_path, "w", encoding="utf-8") as f:
        # "w" = write mode (creates file or overwrites)
        # encoding="utf-8" handles special characters properly
        f.write("\n\n--- Page Break ---\n\n".join(all_text))
        # Join all pages with a separator so you can tell where pages end

    return output_path


def _pdf_to_word(input_path, output_dir, base_name):
    """Convert PDF → DOCX by extracting text and writing to Word"""

    output_path = os.path.join(output_dir, f"{base_name}.docx")

    # Extract text from PDF
    with pdfplumber.open(input_path) as pdf:
        all_text = []
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                all_text.append(text)

    # Write extracted text into a new Word document
    doc = DocxDocument()
    # DocxDocument() with no arguments creates a blank Word document

    for page_text in all_text:
        for line in page_text.split("\n"):
            para = doc.add_paragraph(line)
            # add_paragraph() adds a new paragraph to the Word doc
            para.runs[0].font.size = Pt(11) if para.runs else None
            # Pt(11) sets font size to 11 points
            # para.runs is a list of "runs" (text segments with same formatting)

    doc.save(output_path)
    return output_path


def _txt_to_pdf(input_path, output_dir, base_name):
    """Convert TXT → PDF using reportlab"""

    output_path = os.path.join(output_dir, f"{base_name}.pdf")

    with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
        # errors="ignore" skips any characters that can't be decoded
        lines = f.readlines()
        # readlines() returns a list where each item is one line of the file

    pdf = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    y = height - 50
    pdf.setFont("Helvetica", 11)

    for line in lines:
        line = line.rstrip("\n")
        # rstrip removes trailing newline characters from each line

        if y < 50:
            pdf.showPage()
            y = height - 50
            pdf.setFont("Helvetica", 11)

        pdf.drawString(40, y, line[:100])
        # [:100] limits to 100 chars per line to prevent overflow
        y -= 16

    pdf.save()
    return output_path


def _jpg_to_pdf(input_path, output_dir, base_name):
    """Convert JPG/PNG image → PDF using Pillow + reportlab"""

    output_path = os.path.join(output_dir, f"{base_name}.pdf")

    img = Image.open(input_path)
    # Image.open() reads the image file into memory

    img_width, img_height = img.size
    # .size returns (width, height) in pixels

    # Scale image to fit A4 page width (595 points)
    a4_width, a4_height = A4
    scale = a4_width / img_width
    # Calculate how much to shrink/grow the image to fit the page width

    new_width = a4_width
    new_height = img_height * scale
    # Apply the same scale to height to keep proportions

    pdf = canvas.Canvas(output_path, pagesize=(new_width, new_height))
    pdf.drawImage(input_path, 0, 0, width=new_width, height=new_height)
    # drawImage(path, x, y, width, height) draws the image onto the PDF canvas
    pdf.save()
    return output_path


def _pdf_to_jpg(input_path, output_dir, base_name):
    """Convert first page of PDF → JPG using pdfplumber + Pillow"""

    output_path = os.path.join(output_dir, f"{base_name}.jpg")

    with pdfplumber.open(input_path) as pdf:
        first_page = pdf.pages[0]
        # Get just the first page — index 0

        img = first_page.to_image(resolution=150)
        # to_image() renders the page as a PIL image
        # resolution=150 means 150 DPI (dots per inch) — good quality

        img.save(output_path)
        # Save the rendered page as a JPG file

    return output_path


def _html_to_pdf(input_path, output_dir, base_name):
    """Convert HTML file → PDF by stripping tags and writing to PDF"""

    output_path = os.path.join(output_dir, f"{base_name}.pdf")

    with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    # Strip HTML tags using a simple approach
    import re
    clean = re.sub(r"<[^>]+>", " ", content)
    # re.sub replaces all matches of the pattern with " "
    # r"<[^>]+>" matches anything between < and > (HTML tags)
    # [^>]+ means "one or more characters that are NOT >"

    clean = re.sub(r"\s+", " ", clean).strip()
    # \s+ matches one or more whitespace characters (spaces, tabs, newlines)
    # This collapses multiple spaces into one

    lines = [clean[i:i+90] for i in range(0, len(clean), 90)]
    # Split into chunks of 90 characters each for line wrapping
    # range(0, len, 90) gives: 0, 90, 180, 270...
    # clean[i:i+90] slices 90 characters at a time

    pdf = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    y = height - 50
    pdf.setFont("Helvetica", 11)

    for line in lines:
        if y < 50:
            pdf.showPage()
            y = height - 50
            pdf.setFont("Helvetica", 11)
        pdf.drawString(40, y, line)
        y -= 16

    pdf.save()
    return output_path
